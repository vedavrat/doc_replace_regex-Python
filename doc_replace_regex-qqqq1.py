import re 
import docx
from docx import Document

def doc_replace_regex(obj_document, txt_regex_pat, txt_replace):
    for p in obj_document.paragraphs:
        print("",p," ~ ",end='') 
        if txt_regex_pat.search(p.text):
            inline = p.runs
            # Loop added to work with "runs" (strings) of paragraph
            for i in range(len(inline)):
                # print(f"{i:2}",end='') 
                if txt_regex_pat.search(inline[i].text):
                    text = txt_regex_pat.sub(txt_replace, inline[i].text)
                    inline[i].text = text
        print("",end='\r') 
    for table in obj_document.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_txt_replace_txt_regex_pat(cell, txt_regex_pat, txt_replace)

callback1 = lambda found1: found1.group(0)[:1] + found1.group(0)[1:].lower()

#egex1 = re.compile(r"regex patthern")
regex1 = re.compile(r"\b[A-Z]{2,}\b")

#replace1 = re.compile(r"replace string")
#replace1 = re.compile(r"\U$0")
#replace1 = re.compile(r"\U&")

file1 = 'libreoffice-abiword-_-file1.docx'
file2 = 'libreoffice-abiword-_-file2.docx'
print("    *** From file: ",file1)
print("    *** To file:   ",file2)
doc1 = Document(file1)
doc_replace_regex(doc1, regex1, callback1)
print()
doc1.save(file2)
